from collections.abc import Generator
import os
import tempfile
import io
import markdown
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Any, List, Dict, Tuple
import html
from bs4 import BeautifulSoup, NavigableString, Tag

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

class DocTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        md_content = tool_parameters.get("markdown_content", "")
        title = tool_parameters.get("title", "Document")
        
        if not md_content:
            yield self.create_text_message("No markdown content provided.")
            return
        
        try:
            # 只去除“↓”，保留“•”
            md_content = md_content.replace("↓", "")
            
            doc = self._convert_markdown_to_docx(md_content, title)
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            file_bytes = docx_bytes.getvalue()
            filename = f"{title.replace(' ', '_')}.docx"
            
            yield self.create_text_message(f"Document '{title}' generated successfully")
            yield self.create_blob_message(
                blob=file_bytes,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": filename
                }
            )
        except Exception as e:
            yield self.create_text_message(f"Error converting markdown to DOCX: {str(e)}")

    def _convert_markdown_to_docx(self, md_content: str, title: str) -> Document:
        doc = Document()
        
        # --- 标题部分 ---
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.italic = False
        title_run.font.name = '方正小标宋_GBK'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 转换 Markdown 为 HTML（移除 nl2br 扩展）
        html_content = markdown.markdown(
            md_content,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.codehilite',
                # 'markdown.extensions.nl2br',
                'markdown.extensions.sane_lists'
            ]
        )
        
        soup = BeautifulSoup(html_content, 'html.parser')
        self._process_html_elements(doc, soup)
        return doc

    def _process_html_elements(self, doc: Document, soup: BeautifulSoup) -> None:
        """
        递归遍历所有 HTML 节点，根据标签名称将内容插入到 Word 文档中。
        只去除“↓”，保留“•”。
        """
        for element in soup.children:
            # 跳过 <br> 标签
            if isinstance(element, Tag) and element.name == 'br':
                continue
            
            # 纯文本节点：只去“↓”
            if isinstance(element, NavigableString):
                text = str(element).replace("↓", "").strip()
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.name = '仿宋GB_2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            
            if not isinstance(element, Tag):
                continue
            
            # 处理标题 h1 ~ h6，只去除“↓”
            if element.name == 'h1':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=1)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h2':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=2)
                for run in heading.runs:
                    run.font.name = '楷体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h3':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=3)
                for run in heading.runs:
                    run.font.name = '仿宋GB_2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h4':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=4)
                for run in heading.runs:
                    run.font.name = '仿宋GB_2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h5':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=5)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h6':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=6)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            # 段落 <p>：只去“↓”
            elif element.name == 'p':
                text = element.get_text().replace("↓", "").strip()
                if not text:
                    continue

                if text.startswith('附件'):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)
                    self._add_run_with_formatting(p, element)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)
                    self._add_run_with_formatting(p, element)

            # 列表 <ul> / <ol>：只去“↓”，保留“•”
            elif element.name == 'ul':
                self._add_list(doc, element, is_numbered=False)

            elif element.name == 'ol':
                self._add_list(doc, element, is_numbered=True)

            # 代码块 <pre>：只去“↓”
            elif element.name == 'pre':
                code = element.get_text().replace("↓", "")
                lang = ""
                if element.code and element.code.get('class'):
                    for cls in element.code.get('class'):
                        if cls.startswith('language-'):
                            lang = cls[9:]
                            break
                self._add_code_block(doc, code, lang)

            # 表格 <table>：只去“↓”
            elif element.name == 'table':
                self._add_html_table(doc, element)

            # 分隔线 <hr>
            elif element.name == 'hr':
                doc.add_paragraph('_' * 50)

            # 其他容器标签（例如 <div>、<span> 等），递归处理其子元素
            else:
                self._process_html_elements(doc, element)

    def _add_run_with_formatting(self, paragraph, element):
        """
        在一个段落里插入富文本内容，并保持“只去除↓”的原则。
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).replace("↓", "")
                run = paragraph.add_run(text)

            elif isinstance(child, Tag) and child.name in ['strong', 'b']:
                text = child.get_text().replace("↓", "")
                run = paragraph.add_run(text)
                run.bold = True

            elif isinstance(child, Tag) and child.name in ['em', 'i']:
                text = child.get_text().replace("↓", "")
                run = paragraph.add_run(text)
                run.italic = True

            elif isinstance(child, Tag) and child.name == 'code':
                code_text = child.get_text().replace("↓", "")
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'

            elif isinstance(child, Tag):
                # 递归处理其他标签（<span>、<a>、<div> 等）
                self._add_run_with_formatting(paragraph, child)
            
            # 统一为正文文字格式：仿宋GB_2312，16pt，黑色
            try:
                run.font.name = '仿宋GB_2312'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
            except UnboundLocalError:
                # 如果 run 未定义（极端情况），忽略即可
                pass

    def _add_list(self, doc: Document, list_element: Tag, is_numbered: bool = False):
        """
        处理 HTML 中的 <ul> 或 <ol> 列表，只去“↓”，保留项目符号。
        """
        for item in list_element.find_all('li', recursive=False):
            # 提取 li 文本并只去“↓”
            item_text = "".join(t for t in item.strings).replace("↓", "").strip()
            if not item_text:
                continue

            # 根据是否有序，选择不同的样式
            p = doc.add_paragraph(style='List Number' if is_numbered else 'List Bullet')
            p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run(item_text)
            run.font.name = '仿宋GB_2312'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)

            # 递归处理嵌套的 <ul> / <ol>
            nested_ul = item.find('ul')
            nested_ol = item.find('ol')
            if nested_ul:
                self._add_list(doc, nested_ul, is_numbered=False)
            if nested_ol:
                self._add_list(doc, nested_ol, is_numbered=True)

    def _add_code_block(self, doc: Document, code: str, language: str = ""):
        """
        将 <pre><code>…</code></pre> 转为 Word 中的代码块格式，只去“↓”。
        """
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        if language:
            lang_run = p.add_run(f"Language: {language}\n")
            lang_run.italic = True
            lang_run.font.name = '仿宋GB_2312'
            lang_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
            lang_run.font.size = Pt(16)
            lang_run.font.color.rgb = RGBColor(0, 0, 0)

        code_filtered = code.replace("↓", "")
        code_run = p.add_run(code_filtered)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(0, 0, 0)

        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)

    def _add_html_table(self, doc: Document, table: Tag):
        """
        将 HTML 中的 <table> 转为 Word 表格，只去“↓”。
        """
        rows = table.find_all('tr')
        if not rows:
            return

        # 第一行决定列数
        header_cells = rows[0].find_all(['th', 'td'])
        col_count = len(header_cells)
        doc_table = doc.add_table(rows=len(rows), cols=col_count)
        doc_table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j >= col_count:
                    continue
                # 只去“↓”，保留其他符号
                cell_text = cell.get_text().replace("↓", "")
                doc_table.cell(i, j).text = cell_text

                for para in doc_table.cell(i, j).paragraphs:
                    for run in para.runs:
                        run.font.name = '仿宋GB_2312'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB_2312')
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                # 如果是表头行（i == 0）或单元格标签是 <th>，则加粗
                if i == 0 or cell.name == 'th':
                    for para in doc_table.cell(i, j).paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
